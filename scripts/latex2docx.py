#!/usr/bin/env python3
"""
Convert LaTeX thesis files to Word (.docx) format.

Handles the LaTeX constructs present in the thesis files:
  - Sections: \chapter, \section, \subsection, \subsubsection
  - Formatting: \textbf{}, \textit{}, \emph{}, \texttt{}, \textsc{}
  - Lists: itemize, enumerate
  - Figures with \includegraphics (inserts PDF/PNG images)
  - Tables (tabular environment, simple parsing)
  - Citations: \cite{...}
  - Cross-references: \ref{}, \label{}
  - Math: inline $...$ and display \[...\]
  - Special LaTeX quotes and dashes
  - \input{...} includes
  - \vspace, \hspace
  - \noindent
  - \textit{}, \textbf{} with ~ ties
"""

import re
import os
import sys
from pathlib import Path

# Add the parent dir of scripts/ to path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import traceback


def resolve_image_path(img_rel, tex_dir, base_dir):
    """
    Try to resolve an image path.
    img_rel: the path from \includegraphics{...}
    tex_dir: directory containing the .tex file
    base_dir: project root (src/)
    """
    # Supported extensions
    exts = ['.pdf', '.png', '.jpg', '.jpeg', '.PNG', '.JPG', '.JPEG', '.PDF']

    # Try paths in order:
    candidates = []

    # 1. Relative to the tex file
    candidates.append(os.path.join(tex_dir, img_rel))
    # 2. Relative to project base (src/)
    candidates.append(os.path.join(base_dir, img_rel))
    # 3. Relative to base with src/ prefix
    candidates.append(os.path.join(base_dir, '..', img_rel))
    # 4. Just the filename part in figures dir
    fname = os.path.basename(img_rel)
    candidates.append(os.path.join(base_dir, 'figures', fname))
    # 5. Various subdirs in the tex dir
    for sub in ['figures', 'images', 'img']:
        candidates.append(os.path.join(tex_dir, sub, fname))

    seen = set()
    for cand in candidates:
        norm = os.path.normpath(cand)
        if norm in seen:
            continue
        seen.add(norm)
        for ext in exts:
            path_with_ext = norm + ext if not norm.lower().endswith(ext.lower()) else norm
            if os.path.exists(path_with_ext):
                return path_with_ext

    return None


def strip_latex_commands(text):
    """Remove common LaTeX commands that don't affect content."""
    # Remove \label{...}
    text = re.sub(r'\\label\{[^}]*\}', '', text)
    # Remove \index{...}
    text = re.sub(r'\\index\{[^}]*\}', '', text)
    # Remove \gls{...} and similar
    text = re.sub(r'\\(gls|Gls|GLS|glspl|Glspl|acr|Acr|acrpl)\{([^}]*)\}', r'\2', text)
    return text


def parse_inline_commands(text, doc=None, paragraph=None, base_dir=None, tex_dir=None, in_table=False):
    """
    Parse LaTeX inline commands in a text string and return a list of (text, bold, italic, font, size, is_image, image_path, image_width) tuples.
    This is a simplified version that handles the most common constructs.
    
    Returns either a list of runs (for normal paragraphs) or plain text (for table cells where we want plain text).
    """
    # For table cells, just strip everything to plain text
    if in_table:
        return latex_to_plain_text(text)
    
    # We'll work with the text and build runs
    result = []

    # If no doc, just return plain text
    if doc is None:
        return latex_to_plain_text(text)

    return text  # Return raw text for now; we'll process at a higher level


def latex_to_plain_text(text):
    """Convert LaTeX text to plain text by stripping commands."""
    # Remove leading/trailing whitespace per line
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        if line:
            cleaned_lines.append(line)
    text = ' '.join(cleaned_lines)

    # Remove comments
    text = re.sub(r'(?<!\\)%.*', '', text)

    # Handle \noindent
    text = re.sub(r'\\noindent\s*', '', text)

    # Handle \vspace and \hspace
    text = re.sub(r'\\(v|h)space\{[^}]*\}', '', text)

    # Handle \addcontentsline
    text = re.sub(r'\\addcontentsline\{[^}]*\}\{[^}]*\}\{[^}]*\}', '', text)

    # Handle \label
    text = re.sub(r'\\label\{[^}]*\}', '', text)

    # Handle \ref, \eqref
    text = re.sub(r'\\(ref|eqref|pageref)\{([^}]*)\}', r'[\2]', text)

    # Handle \cite
    text = re.sub(r'\\cite(?:\[[^\]]*\])?\{([^}]*)\}', lambda m: f'[{m.group(1)}]', text)

    # Handle \textbf
    text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)

    # Handle \textit
    text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)

    # Handle \emph
    text = re.sub(r'\\emph\{([^}]*)\}', r'\1', text)

    # Handle \texttt
    text = re.sub(r'\\texttt\{([^}]*)\}', r'\1', text)

    # Handle \textsc
    text = re.sub(r'\\textsc\{([^}]*)\}', r'\1', text)

    # Handle \textsf
    text = re.sub(r'\\textsf\{([^}]*)\}', r'\1', text)

    # Handle \textit{\textbf{...}} or \textbf{\textit{...}}
    while re.search(r'\\(textbf|textit|emph|texttt)\{', text):
        text = re.sub(r'\\(textbf|textit|emph|texttt)\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}', r'\2', text)

    # Handle inline math $...$
    text = re.sub(r'\$([^\$]+)\$', r'\1', text)

    # Handle display math \[...\] or $$...$$
    text = re.sub(r'\\\[([^\]]*)\\\]', r'\1', text)
    text = re.sub(r'\$\$([^$]+)\$\$', r'\1', text)

    # Handle \includegraphics
    text = re.sub(r'\\includegraphics(?:\[[^\]]*\])?\{([^}]*)\}', '[Figura: \1]', text)

    # Handle \\ (newline in LaTeX)
    text = re.sub(r'\\\\\s*', ' ', text)

    # Handle \& (ampersand)
    text = text.replace('\\&', '&')

    # Handle \%
    text = text.replace('\\%', '%')

    # Handle \_
    text = text.replace('\\_', '_')

    # Handle \textbar
    text = text.replace('\\textbar{}', '|')
    text = text.replace('\\textbar', '|')

    # Handle \textbackslash
    text = text.replace('\\textbackslash{}', '\\')
    text = text.replace('\\textbackslash', '\\')

    # Handle \textasciitilde
    text = text.replace('\\textasciitilde{}', '~')
    text = text.replace('\\textasciitilde', '~')

    # Handle \textdegree
    text = text.replace('\\textdegree{}', '°')
    text = text.replace('\\textdegree', '°')

    # Handle \textmu
    text = text.replace('\\textmu{}', 'μ')
    text = text.replace('\\textmu', 'μ')

    # Handle \texttrademark
    text = text.replace('\\texttrademark{}', '™')
    text = text.replace('\\texttrademark', '™')

    # Handle \textregistered
    text = text.replace('\\textregistered{}', '®')
    text = text.replace('\\textregistered', '®')

    # Handle \textcopyright
    text = text.replace('\\textcopyright{}', '©')
    text = text.replace('\\textcopyright', '©')

    # Handle \textellipsis
    text = text.replace('\\textellipsis{}', '…')
    text = text.replace('\\textellipsis', '…')
    text = text.replace('\\ldots', '…')
    text = text.replace('\\dots', '…')

    # Handle \-- (en-dash)
    text = text.replace('\\--', '–')
    # Handle \--- (em-dash)
    text = text.replace('\\---', '—')

    # Handle quotes
    text = text.replace("``", '"')
    text = text.replace("''", '"')
    text = text.replace("`", "'")

    # Handle ties (~) -> space
    text = text.replace('~', ' ')

    # Handle \begin{...} ... \end{...} blocks - just remove the tags
    text = re.sub(r'\\begin\{[^}]*\}', '', text)
    text = re.sub(r'\\end\{[^}]*\}', '', text)

    # Handle \item
    text = re.sub(r'\\item\s*', '', text)

    # Handle \hline
    text = re.sub(r'\\hline\s*', '', text)

    # Handle \input, \include
    text = re.sub(r'\\(?:input|include)\{[^}]*\}', '', text)

    # Handle \resizebox
    text = re.sub(r'\\resizebox\{[^}]*\}\{[^}]*\}\{', '', text)
    # Close brace of resizebox
    text = re.sub(r'\}\s*$', '', text)

    # Handle remaining commands that are just text inside braces
    # Remove any remaining \[command]{...} patterns
    text = re.sub(r'\\[a-zA-Z]+\*?(?:\[[^\]]*\])?\{([^}]*)\}', r'\1', text)

    # Remove any remaining \command (without braces)
    text = re.sub(r'\\[a-zA-Z]+\*?(?:\{[^}]*\}|\[[^\]]*\])*', '', text)

    # Clean up extra spaces
    text = re.sub(r'  +', ' ', text)
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = text.strip()

    return text


def process_text_with_formatting(text):
    """
    Process a text string that may contain LaTeX formatting commands.
    Returns a list of dicts: { 'text': str, 'bold': bool, 'italic': bool, 'tt': bool }
    """
    runs = []
    pos = 0

    # Pattern for commands
    cmd_pattern = re.compile(
        r'\\(textbf|textit|emph|texttt|textsc|textsf)\{(.*?)\}(?:\s|$)'
        r'|\$([^$]+)\$'
        r'|\\cite(?:\[[^\]]*\])?\{([^}]*)\}'
        r'|\\ref\{([^}]*)\}'
        r'|\\eqref\{([^}]*)\}'
        r'|\\gls\{([^}]*)\}'
        r'|``|\'\'|`'
        r'|\\textdegree\{\}|\\textdegree'
        r'|\\textmu\{\}|\\textmu'
        r'|\\textellipsis\{\}|\\textellipsis|\\ldots|\\dots'
        r'|\\--|\\---'
        r'|\\&|\\\\%|\\\\_'
        r'|~'
    )

    # Actually, we need a different approach - recursively process nested braces
    i = 0
    while i < len(text):
        # Check for \command{...} patterns
        m = re.match(r'\\(textbf|textit|emph|texttt|textsc|textsf)\{(.*?)\}(?:\s|$)', text[i:])

        if m:
            cmd = m.group(1)
            inner = m.group(2)
            is_bold = cmd in ('textbf', 'textsc')
            is_italic = cmd in ('textit', 'emph')
            is_tt = cmd == 'texttt'

            # Recursively process inner content for nested formatting
            inner_runs = process_text_with_formatting(inner)
            for r in inner_runs:
                if is_bold:
                    r['bold'] = True
                if is_italic:
                    r['italic'] = True
                if is_tt:
                    r['tt'] = True
            runs.extend(inner_runs)
            i += m.end()
            continue

        # Check for nested command with potential nested braces inside
        m = re.match(r'\\(textbf|textit|emph|texttt)\{(.*)', text[i:])
        if m:
            cmd = m.group(1)
            inner_start = m.start(2)
            # Find matching closing brace
            depth = 1
            j = inner_start
            while j < len(text[i:]) and depth > 0:
                if text[i+j] == '{':
                    depth += 1
                elif text[i+j] == '}':
                    depth -= 1
                j += 1
            if depth == 0:
                inner = text[i+inner_start : i+j-1]
                is_bold = cmd in ('textbf',)
                is_italic = cmd in ('textit', 'emph')
                is_tt = cmd == 'texttt'
                inner_runs = process_text_with_formatting(inner)
                for r in inner_runs:
                    if is_bold:
                        r['bold'] = True
                    if is_italic:
                        r['italic'] = True
                    if is_tt:
                        r['tt'] = True
                runs.extend(inner_runs)
                i += j
                continue

        # Check for inline math $...$
        m = re.match(r'\$([^$]+)\$', text[i:])
        if m:
            runs.append({'text': m.group(1), 'bold': False, 'italic': True, 'tt': False, 'math': True})
            i += m.end()
            continue

        # Check for \cite
        m = re.match(r'\\cite(?:\[[^\]]*\])?\{([^}]*)\}', text[i:])
        if m:
            runs.append({'text': f'[{m.group(1)}]', 'bold': False, 'italic': False, 'tt': False})
            i += m.end()
            continue

        # Check for \ref / \eqref
        m = re.match(r'\\(?:eq)?ref\{([^}]*)\}', text[i:])
        if m:
            runs.append({'text': f'[{m.group(1)}]', 'bold': False, 'italic': False, 'tt': False})
            i += m.end()
            continue

        # Check for \gls etc.
        m = re.match(r'\\(?:gls|Gls|GLS|glspl|acr|Acr|acrpl)\{([^}]*)\}', text[i:])
        if m:
            runs.append({'text': m.group(1), 'bold': False, 'italic': False, 'tt': False})
            i += m.end()
            continue

        # Quotation marks
        if text[i:i+2] == '``':
            runs.append({'text': '"', 'bold': False, 'italic': False, 'tt': False})
            i += 2
            continue
        if text[i:i+2] == "''":
            runs.append({'text': '"', 'bold': False, 'italic': False, 'tt': False})
            i += 2
            continue
        if text[i] == '`' and (i+1 >= len(text) or text[i+1] != '`'):
            runs.append({'text': "'", 'bold': False, 'italic': False, 'tt': False})
            i += 1
            continue

        # Special characters
        if text[i:i+3] == '\\---':
            runs.append({'text': '—', 'bold': False, 'italic': False, 'tt': False})
            i += 3
            continue
        if text[i:i+2] == '\\--':
            runs.append({'text': '–', 'bold': False, 'italic': False, 'tt': False})
            i += 2
            continue
        if text[i:i+2] == '\\&':
            runs.append({'text': '&', 'bold': False, 'italic': False, 'tt': False})
            i += 2
            continue
        if text[i:i+2] == '\\_':
            runs.append({'text': '_', 'bold': False, 'italic': False, 'tt': False})
            i += 2
            continue
        if text[i:i+2] == '\\%':
            runs.append({'text': '%', 'bold': False, 'italic': False, 'tt': False})
            i += 2
            continue
        # \textdegree
        m = re.match(r'\\textdegree\{\}', text[i:])
        if m:
            runs.append({'text': '°', 'bold': False, 'italic': False, 'tt': False})
            i += m.end()
            continue
        m = re.match(r'\\textdegree\b', text[i:])
        if m and (i+m.end() >= len(text) or text[i+m.end()] != '{'):
            runs.append({'text': '°', 'bold': False, 'italic': False, 'tt': False})
            i += m.end()
            continue
        # \textmu
        if text[i:i+10] == '\\textmu{}' or text[i:i+7] == '\\textmu':
            runs.append({'text': 'μ', 'bold': False, 'italic': False, 'tt': False})
            i += (10 if text[i:i+10] == '\\textmu{}' else 7)
            continue
        # \ldots, \dots, \textellipsis
        m = re.match(r'\\(?:ldots|dots|textellipsis)\{\}', text[i:])
        if m:
            runs.append({'text': '…', 'bold': False, 'italic': False, 'tt': False})
            i += m.end()
            continue
        m = re.match(r'\\(?:ldots|dots|textellipsis)\b', text[i:])
        if m:
            runs.append({'text': '…', 'bold': False, 'italic': False, 'tt': False})
            i += m.end()
            continue

        # Tie (~)
        if text[i] == '~':
            runs.append({'text': ' ', 'bold': False, 'italic': False, 'tt': False})
            i += 1
            continue

        # Escape sequences like \textbar, \textbackslash, etc.
        m = re.match(r'\\text(?:bar|backslash|asciitilde|trademark|registered|copyright)\{\}', text[i:])
        if m:
            replacements = {
                'textbar': '|', 'textbackslash': '\\', 'textasciitilde': '~',
                'texttrademark': '™', 'textregistered': '®', 'textcopyright': '©'
            }
            key = m.group(0).replace('\\', '').replace('{}', '')
            runs.append({'text': replacements.get(key, ''), 'bold': False, 'italic': False, 'tt': False})
            i += m.end()
            continue

        # Any remaining \command (skip it)
        if text[i] == '\\' and i+1 < len(text) and text[i+1].isalpha():
            m = re.match(r'\\[a-zA-Z]+\*?', text[i:])
            if m:
                skip_cmd = m.group()
                # Check if it has optional args
                j = i + m.end()
                # Skip optional args [...]
                while j < len(text) and text[j] == '[':
                    depth = 1
                    j += 1
                    while j < len(text) and depth > 0:
                        if text[j] == '[': depth += 1
                        elif text[j] == ']': depth -= 1
                        j += 1
                # Skip required arg {..} but only if we're not already in nested parsing
                if j < len(text) and text[j] == '{':
                    depth = 1
                    j += 1
                    while j < len(text) and depth > 0:
                        if text[j] == '{': depth += 1
                        elif text[j] == '}': depth -= 1
                        j += 1
                i = j
                continue

        # Normal character
        runs.append({'text': text[i], 'bold': False, 'italic': False, 'tt': False})
        i += 1

    return runs


def add_formatted_text(paragraph, runs_data):
    """Add formatted runs to a paragraph."""
    for run_data in runs_data:
        text = run_data.get('text', '')
        if not text:
            continue
        run = paragraph.add_run(text)
        run.bold = run_data.get('bold', False)
        run.italic = run_data.get('italic', False)
        if run_data.get('tt', False):
            run.font.name = 'Courier New'
        # Math mode: italic
        if run_data.get('math', False):
            run.italic = True


def add_heading_text(doc, text, level=1):
    """Add a heading to the document."""
    heading = doc.add_heading(text, level=level)
    return heading


def add_paragraph_text(doc, text, style=None, bold=False, italic=False, alignment=None, space_after=None):
    """Add a paragraph with optional formatting."""
    if not text.strip():
        return None
    para = doc.add_paragraph()
    if style:
        para.style = style
    if alignment:
        para.alignment = alignment
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)

    runs = process_text_with_formatting(text)
    add_formatted_text(para, runs)
    return para


def process_include(tex_path, base_dir, tex_dir):
    """Process \input{...} or \include{...} commands."""
    content = ''
    try:
        with open(tex_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except:
        return ''

    # Find \input{...} commands
    def resolve_include(m):
        path = m.group(1)
        # Try various paths
        candidates = [
            path + '.tex',
            os.path.join(tex_dir, path + '.tex'),
            os.path.join(base_dir, path + '.tex'),
        ]
        for c in candidates:
            c = os.path.normpath(c)
            if os.path.exists(c):
                try:
                    with open(c, 'r', encoding='utf-8') as f2:
                        return f2.read()
                except:
                    pass
        return f'[INCLUDED: {path}]'

    # Handle \input and \include
    content = re.sub(r'\\(?:input|include)\{([^}]*)\}', resolve_include, content)
    return content


def tex_to_docx(tex_path, output_path=None, base_dir=None):
    """Main function to convert a .tex file to .docx."""
    tex_path = os.path.normpath(tex_path)
    if output_path is None:
        output_path = tex_path.rsplit('.', 1)[0] + '.docx'

    tex_dir = os.path.dirname(tex_path)
    if base_dir is None:
        # Try to find src/ directory
        base_dir = os.path.join(os.path.dirname(os.path.dirname(tex_path)), '..', 'src')
        if not os.path.exists(base_dir):
            base_dir = tex_dir
    
    if not os.path.exists(tex_path):
        print(f"Error: {tex_path} not found", file=sys.stderr)
        return False

    print(f"Converting: {tex_path}")
    
    # Read raw content
    try:
        with open(tex_path, 'r', encoding='utf-8') as f:
            raw_content = f.read()
    except Exception as e:
        print(f"  ERROR reading file: {e}", file=sys.stderr)
        return False

    # Resolve \input{} and \include{} commands inline
    def resolve_input(m):
        path = m.group(1)
        # Try with .tex extension
        cand_paths = [
            path + '.tex',
            path if path.endswith('.tex') else path,
        ]
        # Also try relative to tex_dir
        rel_path = os.path.join(tex_dir, path)
        cand_paths.append(rel_path + '.tex')
        cand_paths.append(rel_path)
        # Try relative to base_dir
        base_path = os.path.join(base_dir, path)
        cand_paths.append(base_path + '.tex')
        cand_paths.append(base_path)
        
        for cand in cand_paths:
            cand = os.path.normpath(cand)
            if os.path.exists(cand):
                try:
                    with open(cand, 'r', encoding='utf-8') as f_incl:
                        return f_incl.read()
                except:
                    pass
        return f'[INCLUYE: {path}]'

    content = re.sub(r'\\(?:input|include)\{([^}]*)\}', resolve_input, raw_content)

    # Now process the content
    doc = Document()

    # === Page setup (A4) ===
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    # === Styles ===
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.15

    # Configure heading styles
    for level, (size, color_hex) in enumerate([
        (Pt(16), '1F3864'),  # Heading 1 (chapter)
        (Pt(14), '2E75B6'),  # Heading 2 (section)
        (Pt(12), '2E75B6'),  # Heading 3 (subsection)
        (Pt(11), '333333'),  # Heading 4
    ], start=1):
        try:
            h_style = doc.styles[f'Heading {level}']
            h_style.font.name = 'Arial'
            h_style.font.size = size
            h_style.font.bold = True
            h_style.font.color.rgb = RGBColor.from_string(color_hex)
            h_style.paragraph_format.space_before = Pt(12)
            h_style.paragraph_format.space_after = Pt(6)
        except:
            pass

    # === Remove comments ===
    # Remove full-line comments
    content = re.sub(r'^\s*%.*$', '', content, flags=re.MULTILINE)

    # === Process structure ===
    lines = content.split('\n')
    
    in_list_itemize = False
    in_list_enumerate = False
    list_counter = 0
    in_figure = False
    in_table_env = False
    in_tabular = False
    tabular_content = ''
    figure_caption = ''
    figure_paths = []
    in_math_display = False
    math_content = ''
    
    skip_patterns = [
        r'\\documentclass',
        r'\\usepackage',
        r'\\bibliographystyle',
        r'\\bibliography',
        r'\\addbibresource',
        r'\\begin\{document\}',
        r'\\end\{document\}',
        r'\\makeindex',
        r'\\input\{.*\}',  # Already resolved
        r'\\include\{.*\}',  # Already resolved
    ]

    def process_inline_content(text):
        """Process a line of text with possible inline LaTeX commands."""
        if not text.strip():
            return None
        
        # Remove \label commands
        text = re.sub(r'\\label\{[^}]*\}', '', text)
        
        # Handle \addcontentsline
        text = re.sub(r'\\addcontentsline\{[^}]*\}\{[^}]*\}\{[^}]*\}', '', text)

        # Handle \noindent
        text = re.sub(r'\\noindent\s*', '', text)
        
        # Handle \vspace, \hspace
        text = re.sub(r'\\(v|h)space\{[^}]*\}', '', text)
        text = re.sub(r'\\(v|h)space\*\?\{[^}]*\}', '', text)
        
        # Handle \resizebox{width}!{content} - keep the content
        m = re.search(r'\\resizebox\{[^}]*\}\{[^}]*\}\{', text)
        if m:
            # Find matching closing brace
            start = m.end()
            depth = 1
            j = start
            while j < len(text) and depth > 0:
                if text[j] == '{': depth += 1
                elif text[j] == '}': depth -= 1
                j += 1
            if depth == 0 and j > start:
                inner = text[start:j-1]
                return process_inline_content(inner)
        
        # Handle \small, \large, etc.
        text = re.sub(r'\\(?:small|footnotesize|tiny|large|Large|LARGE|huge|Huge|normalsize)\s*', '', text)
        
        # Handle \centering
        text = re.sub(r'\\centering\s*', '', text)

        # Handle \textbf, \textit, \emph, etc.
        # These will be handled by process_text_with_formatting
        
        return text

    def add_formatted_paragraph(doc, text, bold=False, italic=False, alignment=None, space_before=0, space_after=6):
        """Add a paragraph with formatted text."""
        text = text.strip()
        if not text:
            return
        
        text = process_inline_content(text)
        if text is None or not text.strip():
            return
        
        para = doc.add_paragraph()
        if alignment:
            para.alignment = alignment
        if space_before:
            para.paragraph_format.space_before = Pt(space_before)
        if space_after:
            para.paragraph_format.space_after = Pt(space_after)
        
        if bold and not any(c in text for c in ['\\textbf', '\\textit', '\\emph']):
            # Simple case: entire paragraph is bold
            run = para.add_run(text.strip())
            run.bold = True
            if italic:
                run.italic = True
        else:
            runs = process_text_with_formatting(text)
            add_formatted_text(para, runs)
        
        return para

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip blank lines in certain modes
        if not stripped and not in_tabular and not in_figure:
            # End list if we hit a blank line
            if in_list_itemize:
                in_list_itemize = False
            if in_list_enumerate:
                in_list_enumerate = False
            i += 1
            continue

        # Skip known preamble patterns
        is_skip = False
        for pat in skip_patterns:
            if re.match(pat, stripped):
                is_skip = True
                break
        if is_skip:
            i += 1
            continue

        # === Display math mode ===
        if stripped.startswith('\\[') or stripped.startswith('$$'):
            in_math_display = True
            math_content = ''
            i += 1
            continue
        if in_math_display:
            if stripped.endswith('\\]') or stripped.endswith('$$'):
                # End of display math
                if math_content.strip():
                    para = doc.add_paragraph()
                    run = para.add_run(math_content.strip())
                    run.italic = True
                in_math_display = False
                i += 1
                continue
            math_content += stripped + ' '
            i += 1
            continue

        # === Figure environment ===
        if stripped.startswith('\\begin{figure}'):
            in_figure = True
            figure_caption = ''
            figure_paths = []
            i += 1
            continue
        if in_figure:
            if stripped.startswith('\\end{figure}'):
                # Add the figure and caption
                for fp in figure_paths:
                    actual_path = resolve_image_path(fp, tex_dir, base_dir)
                    if actual_path and os.path.exists(actual_path):
                        try:
                            doc.add_picture(actual_path, width=Inches(5.5))
                        except:
                            para = doc.add_paragraph()
                            run = para.add_run(f'[Imagen: {fp}]')
                            run.italic = True
                    else:
                        para = doc.add_paragraph()
                        run = para.add_run(f'[Imagen no encontrada: {fp}]')
                        run.italic = True
                
                if figure_caption:
                    para = doc.add_paragraph()
                    run = para.add_run(f'Figura: {figure_caption}')
                    run.italic = True
                    run.font.size = Pt(10)
                
                in_figure = False
                i += 1
                continue
            
            # Extract caption
            m = re.match(r'\\caption\{(.*)\}', stripped)
            if m:
                figure_caption = m.group(1)
                i += 1
                continue
            
            # Extract includegraphics
            m = re.search(r'\\includegraphics(?:\[[^\]]*\])?\{(.*?)\}', stripped)
            if m:
                figure_paths.append(m.group(1))
                i += 1
                continue
            
            # Skip other figure content (\centering, \label, etc.)
            i += 1
            continue

        # === Table environment ===
        if stripped.startswith('\\begin{table}'):
            in_table_env = True
            i += 1
            continue
        if in_table_env:
            if stripped.startswith('\\end{table}'):
                in_table_env = False
                i += 1
                continue
            # Pass through to normal processing for tabular inside table
            if stripped.startswith('\\begin{tabular}'):
                pass
            elif stripped.startswith('\\caption'):
                m = re.match(r'\\caption\{(.*)\}', stripped)
                if m:
                    para = doc.add_paragraph()
                    run = para.add_run(f'Tabla: {m.group(1)}')
                    run.italic = True
                    run.font.size = Pt(10)
                i += 1
                continue
            elif stripped.startswith('\\label') or stripped.startswith('\\centering') or not stripped:
                i += 1
                continue

        # === Tabular environment ===
        if stripped.startswith('\\begin{tabular}'):
            in_tabular = True
            tabular_content = ''
            m = re.search(r'\\begin\{tabular\}\{([^}]*)\}', stripped)
            cols_spec = m.group(1) if m else ''
            i += 1
            continue
        
        if in_tabular:
            if stripped.startswith('\\end{tabular}'):
                in_tabular = False
                # Parse and render the tabular content as a Word table
                rows_data = tabular_content.strip().split('\\\\')
                # Clean up each row
                clean_rows = []
                for row_text in rows_data:
                    row_text = row_text.strip()
                    # Remove \hline
                    row_text = re.sub(r'\\hline\s*', '', row_text)
                    if row_text:
                        # Split by &
                        cells = [c.strip() for c in row_text.split('&')]
                        clean_rows.append(cells)
                
                if clean_rows:
                    num_cols = max(len(r) for r in clean_rows)
                    table = doc.add_table(rows=len(clean_rows), cols=num_cols)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    for r_idx, row_cells in enumerate(clean_rows):
                        for c_idx, cell_text in enumerate(row_cells):
                            if c_idx < num_cols:
                                cell = table.cell(r_idx, c_idx)
                                cell_text_plain = latex_to_plain_text(cell_text)
                                # Check if bold (\\textbf in first row usually)
                                is_header = (r_idx == 0) or ('textbf' in cell_text)
                                cell.text = ''
                                para = cell.paragraphs[0]
                                run = para.add_run(cell_text_plain.strip())
                                if is_header:
                                    run.bold = True
                
                i += 1
                continue
            
            # Accumulate tabular content
            # Handle \hline
            if stripped.startswith('\\hline'):
                i += 1
                continue
            # Handle \label inside tabular
            if stripped.startswith('\\label'):
                i += 1
                continue
            tabular_content += line + '\n'
            i += 1
            continue

        # === Itemize / Enumerate ===
        if stripped.startswith('\\begin{itemize}'):
            in_list_itemize = True
            i += 1
            continue
        if stripped.startswith('\\end{itemize}'):
            in_list_itemize = False
            i += 1
            continue
        if stripped.startswith('\\begin{enumerate}'):
            in_list_enumerate = True
            list_counter = 0
            i += 1
            continue
        if stripped.startswith('\\end{enumerate}'):
            in_list_enumerate = False
            i += 1
            continue

        if in_list_itemize:
            if stripped.startswith('\\item'):
                item_text = stripped[5:].strip()
                item_text = process_inline_content(item_text)
                if item_text:
                    para = doc.add_paragraph(style='List Bullet')
                    runs = process_text_with_formatting(item_text)
                    add_formatted_text(para, runs)
                i += 1
                continue
            else:
                # Continuation of last item
                if stripped:
                    # Get the last paragraph and add more text
                    if doc.paragraphs:
                        last_para = doc.paragraphs[-1]
                        if last_para.style.name == 'List Bullet':
                            line_text = process_inline_content(stripped)
                            if line_text:
                                runs = process_text_with_formatting(line_text)
                                add_formatted_text(last_para, runs)
                i += 1
                continue

        if in_list_enumerate:
            if stripped.startswith('\\item'):
                list_counter += 1
                item_text = stripped[5:].strip()
                item_text = process_inline_content(item_text)
                if item_text:
                    # Use List Number style
                    para = doc.add_paragraph(style='List Number')
                    runs = process_text_with_formatting(item_text)
                    add_formatted_text(para, runs)
                i += 1
                continue
            else:
                if stripped:
                    if doc.paragraphs:
                        last_para = doc.paragraphs[-1]
                        if last_para.style.name == 'List Number':
                            line_text = process_inline_content(stripped)
                            if line_text:
                                runs = process_text_with_formatting(line_text)
                                add_formatted_text(last_para, runs)
                i += 1
                continue

        # === Section headings ===
        # \chapter{...} or \chapter*{...}
        m = re.match(r'\\(?:chapter\*?)\{(.*)\}', stripped)
        if m:
            title = m.group(1)
            title = latex_to_plain_text(title)
            if title.strip():
                add_heading_text(doc, title.strip(), level=1)
            i += 1
            continue

        # \section{...}
        m = re.match(r'\\(?:section\*?)\{(.*)\}', stripped)
        if m:
            title = m.group(1)
            title = latex_to_plain_text(title)
            if title.strip():
                add_heading_text(doc, title.strip(), level=2)
            i += 1
            continue

        # \subsection{...}
        m = re.match(r'\\(?:subsection\*?)\{(.*)\}', stripped)
        if m:
            title = m.group(1)
            title = latex_to_plain_text(title)
            if title.strip():
                add_heading_text(doc, title.strip(), level=3)
            i += 1
            continue

        # \subsubsection{...}
        m = re.match(r'\\(?:subsubsection\*?)\{(.*)\}', stripped)
        if m:
            title = m.group(1)
            title = latex_to_plain_text(title)
            if title.strip():
                add_heading_text(doc, title.strip(), level=4)
            i += 1
            continue

        # === Special environments ===
        # \begin{titlepage} ... \end{titlepage}
        if stripped.startswith('\\begin{titlepage}'):
            # Skip until \end{titlepage}
            while i < len(lines) and not lines[i].strip().startswith('\\end{titlepage}'):
                i += 1
            i += 1
            continue

        # \begin{abstract} ... \end{abstract}
        if stripped.startswith('\\begin{abstract}'):
            i += 1
            continue
        if stripped.startswith('\\end{abstract}'):
            i += 1
            continue

        # \begin{center} ... \end{center}
        if stripped.startswith('\\begin{center}'):
            in_center = True
            i += 1
            continue
        if stripped.startswith('\\end{center}'):
            in_center = False
            i += 1
            continue

        # Other \begin{...} / \end{...} to skip
        if re.match(r'\\begin\{[^}]*\}', stripped):
            i += 1
            continue
        if re.match(r'\\end\{[^}]*\}', stripped):
            i += 1
            continue

        # === Plain text paragraphs ===
        if stripped:
            # Check if it's just a command we should skip
            if re.match(r'\\[a-zA-Z]', stripped) and '{' not in stripped and '}' not in stripped:
                # Just a standalone command, like \newpage, \clearpage, etc.
                i += 1
                continue
            
            # Process and add as paragraph
            text = process_inline_content(stripped)
            if text and text.strip():
                # Check for bold in text
                is_bold = bool(re.search(r'\\textbf\{', stripped))
                # Check if it looks like a section-like line (all bold)
                para = add_formatted_paragraph(doc, text, bold=is_bold)
        else:
            # Empty line - add small spacing
            pass

        i += 1

    # Save the document
    try:
        doc.save(output_path)
        print(f"  -> {output_path}")
        return True
    except Exception as e:
        print(f"  ERROR saving {output_path}: {e}", file=sys.stderr)
        traceback.print_exc()
        return False


def find_base_dir(tex_path):
    """Find the project base directory (src/) from a .tex file path."""
    p = os.path.dirname(os.path.abspath(tex_path))
    for _ in range(5):
        if os.path.basename(p) == 'src':
            return p
        parent = os.path.dirname(p)
        if parent == p:
            break
        p = parent
    # Fallback: look for figures/ directory
    p = os.path.dirname(os.path.abspath(tex_path))
    for _ in range(5):
        if os.path.exists(os.path.join(p, 'figures')):
            return p
        parent = os.path.dirname(p)
        if parent == p:
            break
        p = parent
    return None


if __name__ == '__main__':
    # Process all .tex files
    tex_files = sys.argv[1:] if len(sys.argv) > 1 else []
    
    if not tex_files:
        print("Usage: python scripts/latex2docx.py <tex_file1> [tex_file2 ...]")
        print("   Or: python scripts/latex2docx.py --all   (finds all .tex files in src/)")
        sys.exit(1)
    
    if tex_files == ['--all']:
        import glob
        base = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'src')
        patterns = [
            'src/thesis.tex',
            'src/chapters/*/*.tex',
            'src/frontmatter/*.tex',
            'src/backmatter/*.tex',
        ]
        tex_files = []
        for pat in patterns:
            tex_files.extend(glob.glob(os.path.join(os.path.dirname(base), pat)))
        tex_files = sorted(set(tex_files))
    
    success = 0
    failed = 0
    
    for tex_path in tex_files:
        tex_path = os.path.normpath(tex_path)
        if not os.path.exists(tex_path):
            print(f"File not found: {tex_path}", file=sys.stderr)
            failed += 1
            continue
        
        output_path = tex_path.rsplit('.', 1)[0] + '.docx'
        base_dir = find_base_dir(tex_path)
        
        if tex_to_docx(tex_path, output_path, base_dir):
            success += 1
        else:
            failed += 1
    
    print(f"\n{'='*50}")
    print(f"Conversion complete: {success} successful, {failed} failed")
